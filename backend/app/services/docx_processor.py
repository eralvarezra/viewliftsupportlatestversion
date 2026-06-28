import os
from typing import List, Dict, Any
from docx import Document


class DocxProcessor:
    """Process DOCX files and extract text chunks.

    Reads both paragraphs and tables. Groups content into Q&A chunks
    by detecting questions (lines ending with '?' or heading styles),
    keeping each question + its full answer together as one chunk.
    Falls back to paragraph-based chunking if no Q&A structure is found.
    """

    CHARS_PER_TOKEN = 4

    def __init__(self, max_tokens: int = 800):
        if max_tokens <= 0:
            raise ValueError("max_tokens must be a positive integer")
        self.max_tokens = max_tokens

    def _is_question(self, text: str, style_name: str, is_bold: bool) -> bool:
        """Detect if a paragraph is a question/heading."""
        if not text:
            return False
        if text.endswith('?'):
            return True
        if style_name and any(style_name.startswith(h) for h in ('Heading', 'Title')):
            return True
        # Bold short text is likely a section header or question
        if is_bold and len(text) < 250 and not text.endswith('.'):
            return True
        return False

    def _extract_paragraphs(self, doc: Document) -> List[Dict]:
        """Extract all paragraphs with style and bold metadata."""
        items = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            items.append({
                'text': text,
                'style': para.style.name if para.style else '',
                'bold': is_bold,
                'source': 'paragraph',
            })
        return items

    def _extract_tables(self, doc: Document) -> List[Dict]:
        """Extract all table content as text items."""
        items = []
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        cells.append(cell_text)
                # Deduplicate merged cells (Word repeats merged cell text)
                seen = []
                for cell in cells:
                    if not seen or cell != seen[-1]:
                        seen.append(cell)
                cells = seen

                if not cells:
                    continue

                if len(cells) == 1:
                    items.append({
                        'text': cells[0],
                        'style': '',
                        'bold': False,
                        'source': 'table',
                    })
                else:
                    # Multi-cell row: treat first cell as potential question
                    question = cells[0]
                    answer = ' '.join(cells[1:])
                    # If first cell looks like a question, keep them as Q&A
                    if question.endswith('?') or len(question) < 200:
                        items.append({
                            'text': f"Q: {question}\nA: {answer}",
                            'style': 'qa_pair',
                            'bold': False,
                            'source': 'table',
                        })
                    else:
                        items.append({
                            'text': ' | '.join(cells),
                            'style': '',
                            'bold': False,
                            'source': 'table',
                        })
        return items

    def _build_qa_chunks(self, items: List[Dict]) -> List[str]:
        """Group items into Q&A chunks."""
        chunks = []
        current_question = None
        current_answer_parts = []
        intro_parts = []

        for item in items:
            text = item['text']

            # Already formatted Q&A from tables — save as-is
            if item.get('style') == 'qa_pair':
                if current_question:
                    answer = ' '.join(current_answer_parts).strip()
                    chunk = f"Q: {current_question}\nA: {answer}" if answer else current_question
                    chunks.append(chunk)
                    current_question = None
                    current_answer_parts = []
                chunks.append(text)
                continue

            if self._is_question(text, item['style'], item['bold']):
                # Save previous Q&A block
                if current_question:
                    answer = ' '.join(current_answer_parts).strip()
                    chunk = f"Q: {current_question}\nA: {answer}" if answer else current_question
                    chunks.append(chunk)
                elif intro_parts:
                    # Save any intro text before first question
                    chunks.append(' '.join(intro_parts))
                    intro_parts = []
                current_question = text
                current_answer_parts = []
            else:
                if current_question:
                    current_answer_parts.append(text)
                else:
                    intro_parts.append(text)

        # Save last block
        if current_question:
            answer = ' '.join(current_answer_parts).strip()
            chunk = f"Q: {current_question}\nA: {answer}" if answer else current_question
            chunks.append(chunk)
        elif intro_parts:
            chunks.append(' '.join(intro_parts))

        return chunks

    def _split_large_chunk(self, text: str) -> List[str]:
        """Split a chunk that exceeds max_tokens into smaller pieces."""
        max_chars = self.max_tokens * self.CHARS_PER_TOKEN
        if len(text) <= max_chars:
            return [text]

        # Try to split on double newlines first (paragraph boundaries)
        parts = text.split('\n\n')
        result = []
        current = ''
        for part in parts:
            if len(current) + len(part) + 2 <= max_chars:
                current = (current + '\n\n' + part).strip() if current else part
            else:
                if current:
                    result.append(current)
                current = part
        if current:
            result.append(current)
        return result if result else [text[:max_chars]]

    def extract_text(self, filepath: str) -> str:
        """Extract all text from DOCX (paragraphs + tables) as plain string."""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        try:
            doc = Document(filepath)
        except Exception as e:
            raise ValueError(f"Invalid DOCX file: {filepath}") from e

        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return '\n\n'.join(parts)

    def process_docx(self, filepath: str) -> Dict[str, Any]:
        """Process a DOCX file into Q&A-aware chunks.

        Reads paragraphs and tables, groups content by Q&A pairs so each
        chunk contains a complete question with its full answer.
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        try:
            doc = Document(filepath)
        except Exception as e:
            raise ValueError(f"Invalid DOCX file: {filepath}") from e

        paragraph_items = self._extract_paragraphs(doc)
        table_items = self._extract_tables(doc)

        # Merge: paragraphs first, then table Q&A items
        # Table items that are full Q&A pairs go at end as supplemental
        paragraph_qa = [i for i in table_items if i.get('style') != 'qa_pair']
        standalone_qa = [i for i in table_items if i.get('style') == 'qa_pair']

        all_items = paragraph_items + paragraph_qa

        raw_chunks = self._build_qa_chunks(all_items)

        # Add standalone Q&A table items
        for item in standalone_qa:
            raw_chunks.append(item['text'])

        # Split any chunks that are still too large
        final_chunks = []
        for chunk in raw_chunks:
            if chunk.strip():
                final_chunks.extend(self._split_large_chunk(chunk.strip()))

        return {
            'filename': os.path.basename(filepath),
            'chunks': final_chunks,
            'total_chunks': len(final_chunks),
        }

    def chunk_text(self, text: str, max_tokens: int = None) -> List[str]:
        """Split plain text into chunks (legacy compatibility)."""
        if max_tokens is None:
            max_tokens = self.max_tokens
        max_chars = max_tokens * self.CHARS_PER_TOKEN
        words = text.split()
        chunks, current, length = [], [], 0
        for word in words:
            wlen = len(word) + 1
            if length + wlen > max_chars and current:
                chunks.append(' '.join(current))
                current, length = [word], wlen
            else:
                current.append(word)
                length += wlen
        if current:
            chunks.append(' '.join(current))
        return chunks
